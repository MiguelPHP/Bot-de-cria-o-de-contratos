import docx

n5 = str(input("Digite o nome do arquivo: "))
n2 = str(input("Digite seu nome: "))
n3 = str(input("Digite o RG com todos os '.': "))
n4 = str(input("Digite seu CPF com todos os '.': "))
n6 = str(input("Digite o numero do imovel: "))
n7 = str(input("Digite o numero do apartamento/Casa : "))

primeiro_paragrafo = (str('''
    pelo presente instrumento particular, de um lado, como LOCADORA, Antónia Rosângela da Silva, brasileira, casada, nascida em 03/03/1965, registrada sob RG MG 4688431 SSP / MG e CPF 697.168.206-82, residente e domiciliada a Av. Geraldo Campos 347, Bandeirinhas, Betim / MG e, de outro lado, como LOCATÁRIO,''') + n2 + str(''' brasileiro, nascido em 19/05/1999, registrado sob RG ''') + n3 + str(''' e CPF ''') + n4 + str(''' residente Rua Córsega 369, Arquipélago Verde de Betim / MG, resolvem celebrar o presente contrato de locação, o qual deve reger-se pelas seguintes cláusulas e condições: '''))

segundo_paragrafo = (str('''        I.	OBJETO: Constitui objeto do presente contrato de locação de imóvel
 localizado na Av. Gerado Campos,''') + n6 + str(''', Apartamento ''') + n7 + str(''',
 Bandeirinhas, Betim / MG'''))

#terceiro_paragrafo = 

document = docx.Document()

document.add_heading("      Contrato de locação de imóvel residencial", level=1)

document.add_paragraph(primeiro_paragrafo)
document.add_paragraph(segundo_paragrafo)


document.save(n5 +( ".docx"))

print('Contrato criado com sucesso!!!!')
